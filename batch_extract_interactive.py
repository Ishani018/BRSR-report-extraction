"""
Simple PDF to DOCX/JSON converter.
- Processes all PDFs from brsr_reports/final/ one by one
- Converts each to DOCX and JSON with the same filename
- Saves outputs to brsr_reports/outputs/
- Skips PDFs that have already been processed
"""

import sys
import logging
import json
import re
from pathlib import Path
from typing import List
from datetime import datetime

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure project root is on sys.path when running directly
BASE_DIR = Path(__file__).parent
sys.path.insert(0, str(BASE_DIR))

from config.config import LOG_FORMAT, LOG_LEVEL, OUTPUT_BASE_DIR  # noqa: E402
from pipeline.detect_pdf_type import detect_pdf_type, get_pdf_info  # noqa: E402
from pipeline.extract_text import extract_text  # noqa: E402
from pipeline.clean_text import clean_pages  # noqa: E402
from pipeline.section_hierarchy_builder import SectionHierarchyBuilder  # noqa: E402


def _setup_logging() -> None:
    logging.basicConfig(
        level=getattr(logging, LOG_LEVEL, logging.INFO),
        format=LOG_FORMAT,
        handlers=[logging.StreamHandler(sys.stdout)],
    )


def process_pdf_simple(pdf_path: Path, output_dir: Path) -> dict:
    """
    Process a single PDF: extract text and convert to DOCX and JSON.
    Uses the same base filename as the PDF.
    
    Returns dict with status and output file paths.
    """
    logger = logging.getLogger(__name__)
    base_name = pdf_path.stem  # filename without .pdf extension
    
    # Output file paths (same name, different extensions)
    docx_path = output_dir / f"{base_name}.docx"
    json_path = output_dir / f"{base_name}.json"
    
    # Check if already processed
    if docx_path.exists() and json_path.exists():
        logger.info(f"✓ Already processed: {pdf_path.name}")
        return {
            "status": "already_processed",
            "pdf_path": str(pdf_path),
            "docx_path": str(docx_path),
            "json_path": str(json_path),
            "error": None
        }
    
    try:
        logger.info(f"Processing: {pdf_path.name}")
        
        # Step 1: Get PDF info
        pdf_info = get_pdf_info(pdf_path)
        logger.info(f"  Pages: {pdf_info.get('pages', 'N/A')}")
        
        # Step 2: Detect PDF type
        pdf_type, _ = detect_pdf_type(pdf_path)
        logger.info(f"  Type: {pdf_type.upper()}")
        
        # Step 3: Extract text
        pages, extraction_stats = extract_text(pdf_path, pdf_type)
        logger.info(f"  Extracted text from {len(pages)} pages")
        
        # Step 4: Clean text
        cleaned_pages = clean_pages(pages)
        total_chars = sum(p.char_count for p in cleaned_pages)
        logger.info(f"  Cleaned {total_chars:,} characters")
        
        # Step 5: Create DOCX
        logger.info(f"  Creating DOCX: {docx_path.name}")
        doc = Document()
        
        title = doc.add_heading(pdf_path.stem, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Pages: {len(cleaned_pages)}")
        doc.add_page_break()
        
        for page in cleaned_pages:
            page_heading = doc.add_heading(f'Page {page.page_number}', level=2)
            page_heading.style.font.size = Pt(12)
            page_heading.style.font.bold = True
            
            if page.text.strip():
                paragraphs = page.text.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        # Clean text: remove NULL bytes and invalid XML control characters
                        cleaned_para = para_text.strip()
                        # Remove NULL bytes
                        cleaned_para = cleaned_para.replace('\x00', '')
                        # Remove other XML-incompatible control characters (keep only printable chars and newlines/tabs)
                        cleaned_para = ''.join(
                            char for char in cleaned_para 
                            if char.isprintable() or char in ['\n', '\t', '\r']
                        )
                        # Replace any remaining control characters with space
                        cleaned_para = ''.join(
                            char if char.isprintable() or char in ['\n', '\t', '\r'] else ' '
                            for char in cleaned_para
                        )
                        if cleaned_para.strip():
                            para = doc.add_paragraph(cleaned_para.strip())
                            para.style.font.size = Pt(10)
                            para.style.font.name = 'Calibri'
            else:
                doc.add_paragraph("[No text on this page]")
            
            doc.add_paragraph()
        
        doc.save(str(docx_path))
        logger.info(f"  ✓ DOCX saved: {docx_path.name}")
        
        # Step 6: Create JSON with hierarchy
        logger.info(f"  Creating JSON: {json_path.name}")
        full_text = "\n\n".join([p.text for p in cleaned_pages])
        
        hierarchy_builder = SectionHierarchyBuilder(section_type='mdna')
        full_report_hierarchy = hierarchy_builder.build_section_hierarchy(
            text=full_text,
            company=base_name,  # Use filename as company name
            year="",  # Year will be in filename
            section_name="BRSR Report",
            start_page=1,
            end_page=len(cleaned_pages),
            confidence=1.0
        )
        
        hierarchy_builder.export_section_json(full_report_hierarchy, json_path)
        logger.info(f"  ✓ JSON saved: {json_path.name}")
        
        return {
            "status": "success",
            "pdf_path": str(pdf_path),
            "docx_path": str(docx_path),
            "json_path": str(json_path),
            "error": None
        }
        
    except Exception as e:
        logger.error(f"  ✗ Error processing {pdf_path.name}: {e}", exc_info=True)
        return {
            "status": "failed",
            "pdf_path": str(pdf_path),
            "docx_path": None,
            "json_path": None,
            "error": str(e)
        }


def simple_extract() -> None:
    """
    Simple extraction: process all PDFs from final folder one by one.
    Converts each to DOCX and JSON with the same filename.
    """
    _setup_logging()
    logger = logging.getLogger(__name__)

    final_dir = BASE_DIR / "brsr_reports" / "final"
    output_dir = OUTPUT_BASE_DIR
    
    if not final_dir.exists():
        logger.error(f"Final folder not found: {final_dir}")
        logger.error("Create it and put your verified PDFs there.")
        return
    
    # Ensure output directory exists
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Get all PDFs - sort by serial number (numeric) at start of filename
    def get_serial_number(path: Path) -> int:
        """Extract serial number from filename for sorting."""
        match = re.match(r'^(\d+)_', path.name)
        if match:
            return int(match.group(1))
        return 999999  # Files without serial number go to end
    
    all_pdf_files = sorted(final_dir.glob("*.pdf"), key=get_serial_number)
    
    if not all_pdf_files:
        logger.warning(f"No PDF files found in {final_dir}")
        return
    
    # Ask how many to process
    total_pdfs = len(all_pdf_files)
    print("")
    print("HOW MANY PDFs TO PROCESS?")
    print(f"Total PDFs found: {total_pdfs}")
    print(f"Enter number [default: {min(5, total_pdfs)}, or 'all']: ", end="")
    
    try:
        user_input = input().strip().lower()
        if user_input in ("", "all"):
            num_to_process = total_pdfs if user_input == "all" else min(5, total_pdfs)
        elif user_input.isdigit():
            num_to_process = min(int(user_input), total_pdfs)
        else:
            num_to_process = min(5, total_pdfs)
    except (EOFError, KeyboardInterrupt):
        logger.info("\nCancelled by user")
        return
    
    pdf_files = all_pdf_files[:num_to_process]
    
    logger.info("")
    logger.info("=" * 80)
    logger.info(f"Processing {len(pdf_files)} of {total_pdfs} PDF file(s)")
    logger.info(f"Output directory: {output_dir}")
    logger.info("=" * 80)
    logger.info("")
    
    results = []
    
    for i, pdf_path in enumerate(pdf_files, 1):
        logger.info(f"[{i}/{len(pdf_files)}] {pdf_path.name}")
        logger.info("-" * 80)
        
        result = process_pdf_simple(pdf_path, output_dir)
        results.append(result)
        
        logger.info("")
    
    # Summary
    logger.info("=" * 80)
    logger.info("SUMMARY")
    logger.info("=" * 80)
    
    successful = sum(1 for r in results if r["status"] == "success")
    already_done = sum(1 for r in results if r["status"] == "already_processed")
    failed = sum(1 for r in results if r["status"] == "failed")
    
    logger.info(f"Total PDFs: {len(results)}")
    logger.info(f"  ✓ Successfully processed: {successful}")
    logger.info(f"  ⊙ Already processed: {already_done}")
    logger.info(f"  ✗ Failed: {failed}")
    
    if failed > 0:
        logger.info("")
        logger.info("Failed PDFs:")
        for r in results:
            if r["status"] == "failed":
                logger.info(f"  - {Path(r['pdf_path']).name}: {r['error']}")
    
    logger.info("")
    logger.info("Done!")


if __name__ == "__main__":
    simple_extract()

