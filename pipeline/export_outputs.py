"""
Module for exporting processed data to various formats.
"""
import logging
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pipeline.extract_text import PageText
from config.config import OUTPUT_DIR, OUTPUT_BASE_DIR, BRSR_FINANCIAL_YEARS
from pipeline.file_naming import (
    clean_company_name,
    format_brsr_output_filename,
    create_output_path
)

logger = logging.getLogger(__name__)


def create_output_directory(company_name: str, year: str, brsr_mode: bool = False) -> Path:
    """
    Create output directory structure for a company/year.
    
    Args:
        company_name: Name of the company
        year: Report year
        brsr_mode: If True, use BRSR directory structure (outputs/{year}/{CompanyName}/)
        
    Returns:
        Path to the output directory
    """
    if brsr_mode:
        # BRSR structure: outputs/{year}/{CompanyName}/
        cleaned_name = clean_company_name(company_name)
        output_path = OUTPUT_BASE_DIR / year / cleaned_name
    else:
        # Original structure: outputs/{CompanyName}/{year}/
        output_path = OUTPUT_DIR / company_name / year
    
    output_path.mkdir(parents=True, exist_ok=True)
    
    logger.info(f"Created output directory: {output_path}")
    return output_path


def create_brsr_output_directory(
    company_name: str, 
    year: str,
    symbol: Optional[str] = None,
    serial_number: Optional[int] = None
) -> Path:
    """
    Create BRSR-specific output directory structure.
    Uses {serial_number}_{SYMBOL} format like downloads if available.
    
    Args:
        company_name: Name of the company
        year: Financial year (e.g., '2022-23')
        symbol: Optional company symbol
        serial_number: Optional serial number
        
    Returns:
        Path to the output directory (outputs/{year}/{serial_number}_{SYMBOL}/ or outputs/{year}/{CompanyName}/)
    """
    if symbol and serial_number is not None:
        # Use {serial_number}_{SYMBOL} format like downloads
        company_folder = f"{serial_number}_{symbol.upper()}"
        output_path = OUTPUT_BASE_DIR / year / company_folder
        output_path.mkdir(parents=True, exist_ok=True)
        logger.info(f"Created output directory: {output_path}")
        return output_path
    else:
        # Fallback to company name format
        return create_output_directory(company_name, year, brsr_mode=True)


def export_to_docx(
    pages: List[PageText],
    sections: List,
    output_path: Path,
    company_name: str,
    year: str
) -> Path:
    """
    Export extracted text to DOCX format - PAGE BY PAGE for best readability.
    Simpler structure that preserves original document flow.
    
    Args:
        pages: List of PageText objects
        sections: List of Section objects (not used in this version)
        output_path: Directory to save the file
        company_name: Company name
        year: Report year
        
    Returns:
        Path to the created DOCX file
    """
    logger.info(f"Exporting to DOCX: {company_name} {year}")
    
    doc = Document()
    
    # Set document title
    title = doc.add_heading(f'{company_name} - Annual Report {year}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Pages: {len(pages)}")
    doc.add_paragraph(f"Extraction Method: {'OCR' if pages[0].method == 'ocr' else 'Direct Text Extraction'}")
    doc.add_page_break()
    
    # Export page by page for maximum readability and preservation of document flow
    logger.info(f"Adding {len(pages)} pages to document")
    
    for page in pages:
        # Add page marker
        page_heading = doc.add_heading(f'Page {page.page_number}', level=2)
        page_heading.style.font.size = Pt(12)
        page_heading.style.font.bold = True
        
        if page.text.strip():
            # Split into paragraphs - preserve original line breaks for layout
            paragraphs = page.text.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Keep single line breaks within paragraphs for formatting preservation
                    para = doc.add_paragraph(para_text.strip())
                    para.style.font.size = Pt(10)
                    para.style.font.name = 'Calibri'
        else:
            doc.add_paragraph("[No text on this page]")
        
        # Add small spacing between pages
        doc.add_paragraph()
    
    # Save document
    docx_path = output_path / "report.docx"
    doc.save(str(docx_path))
    
    logger.info(f"DOCX saved to: {docx_path}")
    return docx_path


def export_brsr_to_docx(
    pages: List[PageText],
    output_path: Path,
    company_name: str,
    year: str,
    is_standalone: bool = True,
    is_from_annual: bool = False,
    naming_convention: Optional[str] = None,
    symbol: Optional[str] = None,
    serial_number: Optional[int] = None
) -> Path:
    """
    Export BRSR content to DOCX with standardized naming.
    
    Args:
        pages: List of PageText objects
        output_path: Directory to save the file
        company_name: Company name
        year: Financial year
        is_standalone: True if standalone BRSR
        is_from_annual: True if extracted from annual report
        naming_convention: Optional naming convention from Excel
        symbol: Optional company symbol for naming convention
        serial_number: Optional serial number for naming convention
        
    Returns:
        Path to the created DOCX file
    """
    logger.info(f"Exporting BRSR to DOCX: {company_name} {year}")
    
    # Create output directory if needed
    output_path = Path(output_path)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Format filename
    filename = format_brsr_output_filename(
        company_name, year, is_standalone, is_from_annual, 'docx',
        naming_convention=naming_convention, symbol=symbol, serial_number=serial_number
    )
    docx_path = output_path / filename
    
    # Create document
    doc = Document()
    
    # Set document title
    title_text = f'{company_name} - BRSR {year}'
    if is_from_annual:
        title_text += ' (from Annual Report)'
    
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Pages: {len(pages)}")
    doc.add_paragraph(f"Type: {'Standalone BRSR' if is_standalone else 'Embedded BRSR'}")
    doc.add_page_break()
    
    # Export page by page
    for page in pages:
        page_heading = doc.add_heading(f'Page {page.page_number}', level=2)
        page_heading.style.font.size = Pt(12)
        page_heading.style.font.bold = True
        
        if page.text.strip():
            paragraphs = page.text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    para.style.font.size = Pt(10)
                    para.style.font.name = 'Calibri'
        else:
            doc.add_paragraph("[No text on this page]")
        
        doc.add_paragraph()
    
    # Save document
    doc.save(str(docx_path))
    logger.info(f"BRSR DOCX saved to: {docx_path}")
    return docx_path




def export_metadata_to_json(
    pdf_info: Dict[str, Any],
    pdf_type: str,
    pages: List[PageText],
    sections: List,
    output_path: Path,
    company_name: str,
    year: str,
    financial_data: Dict[str, Any] = None
) -> Path:
    """
    Export metadata and processing information to JSON.
    
    Args:
        pdf_info: PDF file information
        pdf_type: Type of PDF (text/scanned)
        pages: List of PageText objects
        sections: List of Section objects
        tables: List of ExtractedTable objects
        output_path: Directory to save the file
        company_name: Company name
        year: Report year
        financial_data: Extracted financial data (optional)
        
    Returns:
        Path to the created JSON file
    """
    logger.info("Exporting metadata to JSON")
    
    metadata = {
        "company": company_name,
        "year": year,
        "processing_date": datetime.now().isoformat(),
        "pdf_info": pdf_info,
        "pdf_type": pdf_type,
        "statistics": {
            "total_pages": len(pages),
            "total_characters": sum(page.char_count for page in pages),
            "total_sections": len(sections),
            "total_tables": len(tables)
        },
        "sections": [
            {
                "title": section.title,
                "level": section.level,
                "start_page": section.start_page,
                "end_page": section.end_page,
                "character_count": len(section.content)
            }
            for section in sections
        ],
        "pages": [
            {
                "page_number": page.page_number,
                "character_count": page.char_count,
                "extraction_method": page.method
            }
            for page in pages
        ]
    }
    
    # Add financial data if available
    if financial_data:
        metadata["financial_data"] = financial_data
    
    json_path = output_path / "metadata.json"
    
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Metadata saved to: {json_path}")
    return json_path


def export_all(
    pdf_info: Dict[str, Any],
    pdf_type: str,
    pages: List[PageText],
    sections: List,
    company_name: str,
    year: str,
    financial_data: Dict[str, Any] = None
) -> Dict[str, Any]:
    """
    Export all data in all formats.
    
    Args:
        pdf_info: PDF file information
        pdf_type: Type of PDF (text/scanned)
        pages: List of PageText objects
        sections: List of Section objects
        company_name: Company name
        year: Report year
        financial_data: Extracted financial data (optional)
        
    Returns:
        Dictionary with paths to all created files
    """
    logger.info(f"Starting export for {company_name} {year}")
    
    # Create output directory
    output_path = create_output_directory(company_name, year)
    
    results = {
        "output_directory": str(output_path),
        "files_created": []
    }
    
    # Export to DOCX
    try:
        docx_path = export_to_docx(pages, sections, output_path, company_name, year)
        results["docx"] = str(docx_path)
        results["files_created"].append(str(docx_path))
    except Exception as e:
        logger.error(f"Error exporting to DOCX: {e}")
        results["docx_error"] = str(e)
    
    # Export metadata to JSON
    try:
        json_path = export_metadata_to_json(
            pdf_info, pdf_type, pages, sections,
            output_path, company_name, year, financial_data
        )
        results["metadata"] = str(json_path)
        results["files_created"].append(str(json_path))
    except Exception as e:
        logger.error(f"Error exporting metadata: {e}")
        results["metadata_error"] = str(e)
    
    logger.info(f"Export completed. Created {len(results['files_created'])} files")
    return results


def create_summary_report(company_results: List[Dict[str, Any]]) -> Path:
    """
    Create a summary report for all processed companies.
    
    Args:
        company_results: List of processing results for each company
        
    Returns:
        Path to the summary report
    """
    summary_path = OUTPUT_DIR / "processing_summary.json"
    
    summary = {
        "processing_date": datetime.now().isoformat(),
        "total_companies": len(company_results),
        "companies": company_results
    }
    
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Summary report saved to: {summary_path}")
    return summary_path
