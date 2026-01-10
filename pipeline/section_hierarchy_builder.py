"""
Hierarchical structure builder for extracted sections.

Detects headings and subheadings using deterministic heuristics and builds
a hierarchical JSON representation of the section content.
"""
import logging
import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, asdict
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


@dataclass
class Heading:
    """Represents a detected heading."""
    text: str
    level: int
    line_index: int
    confidence: float


@dataclass
class ContentBlock:
    """Represents a content block (heading with paragraphs and optional subsections)."""
    heading: str
    level: int
    content: List[str]
    subsections: Optional[List['ContentBlock']] = None


class SectionHierarchyBuilder:
    """
    Builds hierarchical structure from flat section text using deterministic heuristics.
    """
    
    # Common MD&A subsection keywords
    MDNA_KEYWORDS = [
        'overview', 'business overview', 'industry overview',
        'financial performance', 'financial review', 'results of operations',
        'operating results', 'segment analysis', 'business segments',
        'revenue', 'expenses', 'profitability', 'cash flow',
        'liquidity', 'capital resources', 'working capital',
        'risk', 'risk factors', 'risk management',
        'outlook', 'future outlook', 'forward looking',
        'strategy', 'business strategy', 'growth strategy',
        'opportunities', 'challenges', 'critical accounting'
    ]
    
    # Common Letter to Stakeholders keywords
    LETTER_KEYWORDS = [
        'dear stakeholders', 'dear shareholders', 'dear investors',
        'year in review', 'highlights', 'achievements',
        'performance', 'strategic initiatives', 'vision',
        'mission', 'values', 'sustainability', 'governance',
        'looking ahead', 'future', 'thank you', 'acknowledgment'
    ]
    
    def __init__(self, section_type: str = 'mdna'):
        """
        Initialize hierarchy builder.
        
        Args:
            section_type: Type of section ('mdna' or 'letter_to_stakeholders')
        """
        self.section_type = section_type
        self.keywords = (
            self.MDNA_KEYWORDS if section_type == 'mdna' 
            else self.LETTER_KEYWORDS
        )
    
    def is_likely_heading(
        self, 
        line: str, 
        prev_line: Optional[str] = None,
        next_line: Optional[str] = None
    ) -> Tuple[bool, float, int]:
        """
        Determine if a line is likely a heading using deterministic heuristics.
        
        Args:
            line: The line to check
            prev_line: Previous line for context
            next_line: Next line for context
            
        Returns:
            Tuple of (is_heading, confidence, estimated_level)
        """
        line = line.strip()
        
        # Skip empty lines or very short lines
        if len(line) < 3:
            return False, 0.0, 0
        
        # Skip lines that look like page numbers or footers
        if re.match(r'^\d+$', line) or re.match(r'^Page \d+', line, re.IGNORECASE):
            return False, 0.0, 0
        
        confidence = 0.0
        level = 2  # Default to level 2
        
        # Heuristic 1: All uppercase (strong heading indicator)
        if line.isupper() and len(line.split()) >= 2:
            confidence += 0.4
            level = 1
            logger.debug(f"Uppercase heading detected: '{line[:50]}'")
        
        # Heuristic 2: Title case (moderate indicator)
        words = line.split()
        title_case_words = sum(1 for w in words if w and w[0].isupper())
        if title_case_words >= len(words) * 0.7 and len(words) >= 2:
            confidence += 0.3
            if level != 1:
                level = 2
            logger.debug(f"Title case heading detected: '{line[:50]}'")
        
        # Heuristic 3: Matches known section keywords
        line_lower = line.lower()
        for keyword in self.keywords:
            if keyword in line_lower:
                confidence += 0.25
                logger.debug(f"Keyword match '{keyword}' in: '{line[:50]}'")
                break
        
        # Heuristic 4: Short line followed by dense paragraph
        if len(line.split()) <= 8 and next_line and len(next_line.strip().split()) > 15:
            confidence += 0.2
            logger.debug(f"Short line followed by paragraph: '{line[:50]}'")
        
        # Heuristic 5: Ends with colon (common in headings)
        if line.endswith(':'):
            confidence += 0.15
            logger.debug(f"Line ends with colon: '{line[:50]}'")
        
        # Heuristic 6: Numeric prefix (e.g., "1. Introduction", "A. Overview")
        if re.match(r'^[A-Z0-9]+[\.\)]\s+[A-Z]', line):
            confidence += 0.25
            level = 2 if not line.isupper() else 1
            logger.debug(f"Numeric/letter prefix detected: '{line[:50]}'")
        
        # Heuristic 7: Standalone line (blank lines before and after)
        if prev_line and next_line:
            if len(prev_line.strip()) == 0 and len(line.split()) <= 10:
                confidence += 0.15
        
        # Adjust level based on length and formatting
        if confidence >= 0.5:
            if len(line.split()) <= 4 and line.isupper():
                level = 1  # Very short, uppercase = top-level heading
            elif len(line.split()) <= 6:
                level = 2  # Medium-length = second-level
            else:
                level = 3  # Longer = third-level
        
        is_heading = confidence >= 0.5
        
        if is_heading:
            logger.info(
                f"Heading detected (confidence: {confidence:.2f}, level: {level}): "
                f"'{line[:60]}...'" if len(line) > 60 else f"'{line}'"
            )
        
        return is_heading, confidence, level
    
    def detect_headings(self, text: str) -> List[Heading]:
        """
        Detect all headings in the text.
        
        Args:
            text: Full section text
            
        Returns:
            List of detected Heading objects
        """
        lines = text.split('\n')
        headings = []
        
        logger.info(f"Analyzing {len(lines)} lines for heading detection...")
        
        for i, line in enumerate(lines):
            prev_line = lines[i-1] if i > 0 else None
            next_line = lines[i+1] if i < len(lines) - 1 else None
            
            is_heading, confidence, level = self.is_likely_heading(
                line, prev_line, next_line
            )
            
            if is_heading:
                headings.append(Heading(
                    text=line.strip(),
                    level=level,
                    line_index=i,
                    confidence=confidence
                ))
        
        logger.info(f"Detected {len(headings)} headings")
        for h in headings:
            logger.debug(
                f"  L{h.level} (conf={h.confidence:.2f}): {h.text[:50]}"
            )
        
        return headings
    
    def build_hierarchy(
        self, 
        text: str, 
        headings: List[Heading]
    ) -> List[Dict[str, Any]]:
        """
        Build hierarchical structure from detected headings and text.
        
        Args:
            text: Full section text
            headings: List of detected headings
            
        Returns:
            List of content blocks representing the hierarchy
        """
        if not headings:
            # No headings detected, treat entire text as single block
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            return [{
                "heading": "Content",
                "level": 1,
                "content": paragraphs
            }]
        
        lines = text.split('\n')
        structure = []
        
        # Sort headings by line index to process in order
        headings_sorted = sorted(headings, key=lambda h: h.line_index)
        
        for i, heading in enumerate(headings_sorted):
            # Determine end of this section (start of next heading or end of text)
            start_line = heading.line_index + 1
            end_line = (
                headings_sorted[i + 1].line_index 
                if i + 1 < len(headings_sorted) 
                else len(lines)
            )
            
            # Extract content between this heading and next
            section_lines = lines[start_line:end_line]
            section_text = '\n'.join(section_lines).strip()
            
            # Split into paragraphs (split on double newlines or empty lines)
            paragraphs = []
            current_para = []
            
            for line in section_lines:
                line = line.strip()
                if not line:
                    if current_para:
                        paragraphs.append(' '.join(current_para))
                        current_para = []
                else:
                    current_para.append(line)
            
            if current_para:
                paragraphs.append(' '.join(current_para))
            
            # Filter out very short paragraphs (likely artifacts)
            paragraphs = [p for p in paragraphs if len(p) > 20]
            
            block = {
                "heading": heading.text,
                "level": heading.level,
                "content": paragraphs
            }
            
            structure.append(block)
        
        # Nest subsections based on heading levels
        nested_structure = self._nest_subsections(structure)
        
        logger.info(f"Built hierarchy with {len(nested_structure)} top-level sections")
        
        return nested_structure
    
    def _nest_subsections(
        self, 
        flat_structure: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Convert flat list of sections into nested hierarchy based on heading levels.
        
        Args:
            flat_structure: Flat list of content blocks
            
        Returns:
            Nested hierarchy
        """
        if not flat_structure:
            return []
        
        nested = []
        stack = []  # Stack to track parent sections at each level
        
        for block in flat_structure:
            level = block['level']
            
            # Pop from stack until we find the correct parent level
            while stack and stack[-1]['level'] >= level:
                stack.pop()
            
            if not stack:
                # Top-level section
                nested.append(block)
                stack.append(block)
            else:
                # Subsection - add to parent
                parent = stack[-1]
                if 'subsections' not in parent:
                    parent['subsections'] = []
                parent['subsections'].append(block)
                stack.append(block)
        
        return nested
    
    def build_section_hierarchy(
        self,
        text: str,
        company: str,
        year: str,
        section_name: str,
        start_page: int,
        end_page: int,
        confidence: float
    ) -> Dict[str, Any]:
        """
        Build complete hierarchical representation of a section.
        
        Args:
            text: Full section text
            company: Company name
            year: Report year
            section_name: Section name (e.g., "Management Discussion and Analysis")
            start_page: Section start page
            end_page: Section end page
            confidence: Boundary detection confidence
            
        Returns:
            Complete JSON structure with metadata and hierarchy
        """
        logger.info(f"Building hierarchy for {section_name}...")
        
        # Detect headings
        headings = self.detect_headings(text)
        
        # Build hierarchy
        structure = self.build_hierarchy(text, headings)
        
        # Assemble final JSON
        result = {
            "company": company,
            "year": year,
            "section": section_name,
            "start_page": start_page,
            "end_page": end_page,
            "confidence": confidence,
            "structure": structure,
            "metadata": {
                "total_headings": len(headings),
                "heading_levels": list(set(h.level for h in headings)),
                "character_count": len(text),
                "paragraph_count": sum(
                    len(block.get('content', [])) 
                    for block in self._flatten_structure(structure)
                )
            }
        }
        
        logger.info(
            f"Hierarchy built: {len(headings)} headings, "
            f"{result['metadata']['paragraph_count']} paragraphs"
        )
        
        return result
    
    def _flatten_structure(
        self, 
        structure: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """Recursively flatten nested structure for counting."""
        result = []
        for block in structure:
            result.append(block)
            if 'subsections' in block:
                result.extend(self._flatten_structure(block['subsections']))
        return result
    
    def export_section_json(
        self,
        hierarchy: Dict[str, Any],
        output_path: Path
    ) -> Path:
        """
        Export section hierarchy to JSON file.
        
        Args:
            hierarchy: Hierarchical structure from build_section_hierarchy()
            output_path: Path to save JSON file
            
        Returns:
            Path to created JSON file
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(hierarchy, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Section JSON exported to: {output_path}")
        return output_path
    
    def export_section_docx(
        self,
        hierarchy: Dict[str, Any],
        output_path: Path
    ) -> Path:
        """
        Export section hierarchy to DOCX with proper heading styles.
        
        Args:
            hierarchy: Hierarchical structure from build_section_hierarchy()
            output_path: Path to save DOCX file
            
        Returns:
            Path to created DOCX file
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(
            f"{hierarchy['company']} - {hierarchy['section']} ({hierarchy['year']})",
            level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(
            f"Pages {hierarchy['start_page']}-{hierarchy['end_page']} | "
            f"Confidence: {hierarchy['confidence']:.0%}"
        ).italic = True
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Add content with hierarchy
        self._add_structure_to_doc(doc, hierarchy['structure'])
        
        # Save
        doc.save(str(output_path))
        
        logger.info(f"Section DOCX exported to: {output_path}")
        return output_path
    
    def _add_structure_to_doc(
        self,
        doc: Document,
        structure: List[Dict[str, Any]],
        base_level: int = 1
    ) -> None:
        """
        Recursively add structured content to DOCX document.
        
        Args:
            doc: Document object
            structure: Hierarchical structure
            base_level: Base heading level for this recursion depth
        """
        for block in structure:
            level = min(block['level'] + base_level - 1, 9)  # DOCX supports up to level 9
            
            # Add heading
            doc.add_heading(block['heading'], level=level)
            
            # Add content paragraphs
            for paragraph_text in block.get('content', []):
                if paragraph_text.strip():
                    para = doc.add_paragraph(paragraph_text)
                    # Add slight spacing between paragraphs
                    para.paragraph_format.space_after = Pt(6)
            
            # Recursively add subsections
            if 'subsections' in block and block['subsections']:
                self._add_structure_to_doc(
                    doc, 
                    block['subsections'], 
                    base_level + 1
                )


def build_section_hierarchy(
    text: str,
    company: str,
    year: str,
    section_name: str,
    section_type: str,
    start_page: int,
    end_page: int,
    confidence: float
) -> Dict[str, Any]:
    """
    Convenience function to build section hierarchy.
    
    Args:
        text: Full section text
        company: Company name
        year: Report year
        section_name: Section name
        section_type: Section type ('mdna' or 'letter_to_stakeholders')
        start_page: Start page number
        end_page: End page number
        confidence: Detection confidence
        
    Returns:
        Hierarchical JSON structure
    """
    builder = SectionHierarchyBuilder(section_type=section_type)
    return builder.build_section_hierarchy(
        text, company, year, section_name, start_page, end_page, confidence
    )


def export_section_json(
    hierarchy: Dict[str, Any],
    output_path: Path
) -> Path:
    """
    Convenience function to export section JSON.
    
    Args:
        hierarchy: Hierarchical structure
        output_path: Output file path
        
    Returns:
        Path to created file
    """
    builder = SectionHierarchyBuilder()
    return builder.export_section_json(hierarchy, output_path)


def export_section_docx(
    hierarchy: Dict[str, Any],
    output_path: Path
) -> Path:
    """
    Convenience function to export section DOCX.
    
    Args:
        hierarchy: Hierarchical structure
        output_path: Output file path
        
    Returns:
        Path to created file
    """
    builder = SectionHierarchyBuilder()
    return builder.export_section_docx(hierarchy, output_path)
